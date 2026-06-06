from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from app.services.preprocessing.cleaner import normalize_whitespace
from app.services.preprocessing.table_converter import table_to_records


def extract_docx_blocks(path: Path, file_name: str) -> list[dict]:
    try:
        doc = DocxDocument(path)
    except Exception as exc:  # pragma: no cover - library error surface
        raise ValueError(f"Failed to extract DOCX file: {file_name}") from exc

    blocks: list[dict] = []
    for paragraph in doc.paragraphs:
        cleaned = normalize_whitespace(paragraph.text)
        if cleaned:
            blocks.append(
                {
                    "text": cleaned,
                    "content_type": "text",
                    "page_number": -1,
                    "table_index": -1,
                    "row_index": -1,
                }
            )

    for table_index, table in enumerate(doc.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        for row_index, record in enumerate(table_to_records(f"Table {table_index}", file_name, rows), start=1):
            blocks.append(
                {
                    "text": normalize_whitespace(record),
                    "content_type": "table",
                    "page_number": -1,
                    "table_index": table_index,
                    "row_index": row_index,
                }
            )
    return blocks
